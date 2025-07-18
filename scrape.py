import re
from pathlib import Path
from bs4 import BeautifulSoup
from fpdf import FPDF
import docx
import argparse

def add_section():
    first = input('What is the number of the first video? ')
    last = input('What is the number of the last video? ')
    for i in range(int(first), int(last)+1):
        with open(Path('sources')/ f"{i}.html", 'w') as touch:
            touch.write('')

def add_review():
    test = input('What test review number is this (will overwrite if file already exists)? ')
    with open(Path('sources')/ f"Practice Test {test}.html", 'w') as touch:
            touch.write('')

def clean_output():
    # clean output folder
    folders = list(Path('output/docx').iterdir()) + list(Path('output/pdf').iterdir())
    for folder in folders:
        for file in folder.iterdir():
            file.unlink()
        folder.rmdir()

    # remove .DS_Store files
    items = Path.cwd().rglob('*')
    for item in items:
        if str(item).find('DS_Store') != -1:
            item.unlink()

def export_word(extract):
    doc = docx.Document()
    section = doc.sections[0]
    section.page_width = docx.shared.Inches(5.5)
    section.page_height = docx.shared.Inches(7.5)
    section.top_margin = docx.shared.Inches(0.5)
    section.bottom_margin = docx.shared.Inches(0.5)
    section.left_margin = docx.shared.Inches(0.5)
    section.right_margin = docx.shared.Inches(0.5)

    doc.add_heading(extract['Title'], 1)

    doc_para = doc.add_paragraph(extract['Body'])
    # doc_para.paragraph_format.space_after = 1
    # doc_run = doc_para.add_run(body)
    # doc_run.font.name = 'Roboto'

    folder = Path(f"output/docx/{extract['Section']}")
    folder.mkdir(exist_ok=True)
    doc.save(folder / f"{extract['Filename']}.docx")

def export_test_word(extract):
    doc = docx.Document()
    section = doc.sections[0]
    section.page_width = docx.shared.Inches(8.5)
    section.page_height = docx.shared.Inches(11)
    section.top_margin = docx.shared.Inches(0.5)
    section.bottom_margin = docx.shared.Inches(0.5)
    section.left_margin = docx.shared.Inches(0.5)
    section.right_margin = docx.shared.Inches(0.5)
    
    # Cover page to write overall and domain-breakdown scores
    doc_para = doc.add_heading(filename, 0)
    table = doc.add_table(rows=1, cols=2)
    row = table.add_row().cells
    row[0].text = 'Date: '
    # CUSTOMIZE: Domain Names or Learning Topics
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    row = table.add_row().cells
    row[0].text = 'Overall: '
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    row = table.add_row().cells
    row[0].text = 'D1 - Domain 1 Name: '
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    row = table.add_row().cells
    row[0].text = 'D2 - Domain 2 Name: '
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    row = table.add_row().cells
    row[0].text = 'D3 - Domain 3 Name: '
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    row = table.add_row().cells
    row[0].text = 'D4 - Domain 4 Name: '
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    row = table.add_row().cells
    row[0].text = 'D5 - Domain 5 Name: '
    row[0].paragraphs[0].alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER
    row[1].text = '________________________________________________________'
    doc.add_page_break()

    # Question: Correct or not, Options, Explanation
    for i, q in enumerate(extract['Body']):
        doc.add_heading(f'Question {i+1}: {q["Correct"]}', 4)
        
        doc_para = doc.add_paragraph(q['Question'])
        doc_para.paragraph_format.space_after = 1
        
        for option in q['Options']:
            doc.add_paragraph(option, style='List Bullet 2')
            
        doc.add_paragraph(f'Explanation: {q["Explanation"]}')

    folder = Path(f"output/docx/{extract['Section']}")
    folder.mkdir(exist_ok=True)
    doc.save(folder / f"{extract['Filename']}.docx")

def export_pdf(extract):
    pdf = FPDF(format='letter', unit='in')
    pdf.add_page()
    pdf.add_font(
        family="HelveticaNeue",
        style="", # regular,
        # CUSTOMIZE: This a MacOS font path, change font, or change path for Windows
        fname= "/System/Library/Fonts/HelveticaNeue.ttc" 
    )
    # CUSTOMIZE: Set font based on name given above
    pdf.set_font("HelveticaNeue", size = 12)
    pdf.set_margins(1, 1)
    effective_page_width = pdf.w - 2*pdf.l_margin
    pdf.multi_cell(
        w=effective_page_width, 
        h=0.22, 
        text= extract['Title'] + '\n\n' + extract['Body'],
        align='J'
        )
    
    folder = Path(f"output/pdf/{extract['Section']}")
    folder.mkdir(exist_ok=True)
    pdf.output(folder / f'{filename}.pdf')

def export_test_pdf(extract):
    pdf = FPDF(format='letter', unit='in')
    pdf.add_page()
    pdf.add_font(
        family="HelveticaNeue",
        style="", # regular,
        # CUSTOMIZE: This a MacOS font path, change font, or change path for Windows
        fname= "/System/Library/Fonts/HelveticaNeue.ttc"
    )
    # CUSTOMIZE: Set font based on name given above
    pdf.set_font("HelveticaNeue", size = 10)
    pdf.set_margins(1, 1)
    effective_page_width = pdf.w - 2*pdf.l_margin
    front_matter = f"{extract['Title']}\n"
    front_matter += """
    Date: ________________________________________________________
    Overall: ________________________________________________________
    D1 - Domain 1 Name: ________________________________________________________
    D2 - Domain 2 Name: ________________________________________________________
    D3 - Domain 3 Name: ________________________________________________________
    D4 - Domain 4 Name: ________________________________________________________
    D5 - Domain 5 Name: ________________________________________________________\n"""
    body = ''
    for i, q in enumerate(extract['Body']):
        body += f"Question {i+1}: {q['Correct']}\n"
        body += f"{q['Question']}\n"
        for choice in q['Options']:
            body += f"      - {choice}\n"
        body += f"{q['Explanation']}\n\n"
    pdf.multi_cell(
        w=effective_page_width, 
        h=0.22, 
        text=front_matter + '\n' + body,
        align='J'
        )
    
    folder = Path(f"output/pdf/{extract['Section']}")
    folder.mkdir(exist_ok=True)
    pdf.output(folder / f'{filename}.pdf')


parser = argparse.ArgumentParser()
parser.add_argument('-c', '--clean', action='store_true', required=False, help='Empty all output')
parser.add_argument('-p', '--prep', action='store_true', required=False, help='Create blank html files for videos')
parser.add_argument('-a', '--add', action='store_true', required=False, help='Add a section of videos')
parser.add_argument('-at', '--addTest', action='store_true', required=False, help='Add a test review page to parse')
parser.add_argument('-docx', '--wordOnly', action='store_true', required=False, help='Only export word files')
parser.add_argument('-pdf', '--pdfOnly', action='store_true', required=False, help='Only export pdf files')
parser.add_argument('-videos', '--transcripts', action='store_true', required=False, help='Only parse video files for transcripts')
parser.add_argument('-tests', '--reviews', action='store_true', required=False, help='Only parse practice test files to review')
args = parser.parse_args()

if args.clean:
    clean_output()

elif args.prep:
    clean_output()
    for item in Path('sources').iterdir():
        item.unlink()

    # create blank html files
    add_section()

elif args.add:
    add_section()

elif args.addTest:
    add_review()

else:
    for source in Path('sources').iterdir():
        if str(source).find('DS_Store') != -1:
            continue
        # Parsing video titles and transcripts
        if source.stem.isnumeric() and not args.reviews:
            try:
                with open(source, 'r') as lesson:
                    soup = BeautifulSoup(lesson, 'html.parser')
                    content = soup.find('div', attrs={'data-purpose':'transcript-panel'})

                    # get video title
                    video_block = soup.find('div', attrs={'data-purpose': 'curriculum-item-viewer-content'})
                    title = video_block.find('section')['aria-label']
            except:
                print(f"EXCEPTION CAUGHT: Error occurred reading {source}\n")
            else:
                # get transcript lines
                lines = content.find_all('span')
                body = ' '.join(
                    list(
                        map(lambda x: re.sub(
                            '^\s+ | \s+$', '', re.sub(
                                '\s+', ' ', x.get_text())
                        ), lines)
                    )
                )

                # extract a filename from the video title
                # Parsing from example: Section 2: Section Name, Lecture 5: Video Name (OBJ. 1.1)
                # CUSTOMIZE: Parsing of video title into a output file name
                filename = title
                if 'OBJ' in filename:
                    filename = filename[ : filename.find('(OBJ')-1 ]
                filename = filename[ filename.find('Lecture') : ] 
                filename = filename.replace(':', ' -').replace('/', ', ')
                
                # extract the section/folder name from the video title
                # CUSTOMIZE: Parsing of video title to get section, making folder w same name to enclose lectures in the same section
                # CUSTOMIZE Option 2: Remove section folders completely
                section = title[ : title.find(', Lecture')] 
                section = section.replace(':', ' -').replace('/', ', ')
                
                extract = {
                    'Section': section,
                    'Filename': filename,
                    'Title': title,
                    'Body': body
                }

                # create and export to docx
                if not args.pdfOnly:
                    export_word(extract)

                # create and export to pdf
                if not args.wordOnly:
                    export_pdf(extract)

        # Parsing Practice Tests
        elif not source.stem.isnumeric() and not args.transcripts:
            try:
                with open(source, 'r') as test:
                    soup = BeautifulSoup(test, 'html.parser')
                    content = soup.find('div', class_='quiz-page-content')
                    
                    filename = Path(test.name).stem
            except:
                print(f"EXCEPTION CAUGHT: Error occurred reading {source}\n")
            else:
                questions = content.find_all('div', class_='result-pane--question-result-pane--sIcOh')

                qs = []
                for question in questions:
                    q = {}
                    # find id = question-prompt, get inner text, 
                    # replace blocks of whitespace with 1 space, 
                    # replace whitespace at start and end of line with nothing
                    q['Question'] = re.sub('^\s{1}|\s{1}$', '', re.sub('\s+', ' ', question.find(id = 'question-prompt').get_text()))
                    # find all id = answer-text, get inner text for each,
                    # use map to 
                        # replace blocks of whitespace for each with 1 space and
                        # replace whitespace at start and end of line with nothing
                    # convert to list
                    q['Options'] = []
                    for choice in question.find_all('div', attrs={'data-purpose': 'answer'}):
                        correct = choice.find('span', attrs = {'data-purpose': 'answer-result-header-user-label'})
                        value = re.sub("^\s{1}|\s{1}$", "", (re.sub("\s+", " ", choice.find('div', attrs = {'data-purpose': 'answer-body'}).get_text())))
                        if correct == None:
                            q['Options'].append(value)
                        else:
                            correct = re.sub("^\s{1}|\s{1}$", "", (re.sub("\s+", " ", correct.get_text())))
                            q['Options'].append(f'{value} - {correct}')
                    # get text that says if this question was right (by unique tag attribute)
                    q['Correct'] = question.find('span', attrs={'data-purpose': 'question-result-header-status-label'}).get_text()
                    # get the id = overall-explanation, get inner text, replace whitespaces like before
                    # remove troubleshoot message at end
                    # CUSTOMIZE: Remove if no explanation, Remove excess info standard across explanation sections
                    q['Explanation'] = re.sub(' For support or reporting issues, include Question ID:.*$', '', re.sub('^\s{1}|\s{1}$', '', re.sub('\s+', ' ', question.find(id = 'overall-explanation').get_text())))
                    # add dict to list of questions
                    qs.append(q)

                extract = {
                    'Section': 'Practice Tests',
                    'Filename': filename,
                    'Title': filename,
                    'Body': qs
                }

                # create and export to docx
                if not args.pdfOnly:
                    export_test_word(extract)

                # create and export to pdf
                if not args.wordOnly:
                    export_test_pdf(extract)
